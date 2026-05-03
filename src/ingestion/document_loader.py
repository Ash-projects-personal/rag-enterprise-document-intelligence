"""
Multi-format document loader supporting PDF, DOCX, and HTML.
Handles 14,000+ enterprise documents with metadata extraction.
"""
import os
from pathlib import Path
from typing import List, Dict, Any
from dataclasses import dataclass


@dataclass
class Document:
    content: str
    metadata: Dict[str, Any]
    doc_id: str


class EnterpriseDocumentLoader:
    """Loads PDF, DOCX, and HTML documents from a directory or list of paths."""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".html", ".htm", ".txt"}

    def __init__(self, source_dir: str = None):
        self.source_dir = source_dir

    def load_pdf(self, path: str) -> Document:
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                text = "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                )
            return Document(
                content=text.strip(),
                metadata={"source": path, "type": "pdf", "pages": len(pdf.pages)},
                doc_id=Path(path).stem,
            )
        except ImportError:
            raise ImportError("Install pdfplumber: pip install pdfplumber")

    def load_docx(self, path: str) -> Document:
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return Document(
                content=text,
                metadata={"source": path, "type": "docx"},
                doc_id=Path(path).stem,
            )
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")

    def load_html(self, path: str) -> Document:
        from html.parser import HTMLParser

        class TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.text_parts = []
                self._skip_tags = {"script", "style"}
                self._current_tag = None

            def handle_starttag(self, tag, attrs):
                self._current_tag = tag

            def handle_data(self, data):
                if self._current_tag not in self._skip_tags:
                    stripped = data.strip()
                    if stripped:
                        self.text_parts.append(stripped)

        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            raw = f.read()
        extractor = TextExtractor()
        extractor.feed(raw)
        text = " ".join(extractor.text_parts)
        return Document(
            content=text,
            metadata={"source": path, "type": "html"},
            doc_id=Path(path).stem,
        )

    def load_file(self, path: str) -> Document:
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            return self.load_pdf(path)
        elif ext in {".docx", ".doc"}:
            return self.load_docx(path)
        elif ext in {".html", ".htm"}:
            return self.load_html(path)
        elif ext == ".txt":
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return Document(
                content=content,
                metadata={"source": path, "type": "txt"},
                doc_id=Path(path).stem,
            )
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def load_directory(self, directory: str = None) -> List[Document]:
        target = directory or self.source_dir
        if not target:
            raise ValueError("No source directory specified.")
        docs = []
        for root, _, files in os.walk(target):
            for fname in files:
                fpath = os.path.join(root, fname)
                if Path(fpath).suffix.lower() in self.SUPPORTED_EXTENSIONS:
                    try:
                        docs.append(self.load_file(fpath))
                    except Exception as e:
                        print(f"[WARN] Skipping {fpath}: {e}")
        print(f"[INFO] Loaded {len(docs)} documents from {target}")
        return docs
